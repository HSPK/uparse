import mimetypes
import pathlib
import re
import uuid
import xml.etree.ElementTree as ET

import requests
from docx import Document as DocxDocument
from docx.document import Document as DocumentObject
from loguru import logger

from uparse.storage import get_storage
from uparse.types import Chunk, Document
from uparse.utils import convert_to, csv_dumps, encode_image_to_base64

from .._base import BaseParser


class WordParser(BaseParser):
    allowed_extensions = [".docx", ".doc"]

    def __init__(self, uri: str | pathlib.Path, output_dir: str = "outputs", **kwargs):
        super().__init__(uri)
        if self.uri.endswith(".doc"):
            self.uri = convert_to(self.uri, format="docx")
        self.output_dir = pathlib.Path(output_dir) / pathlib.Path(self.uri).stem

    def parse(self) -> Document:
        path = self.get_file_path(self.uri)
        return self.parse_docx(path)

    def _extract_images_from_docx(self, doc: DocumentObject):
        image_dir = self.output_dir / "images"
        image_count = 0
        image_map = {}

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                if rel.is_external:
                    url = rel.reltype
                    response = requests.get(url, stream=True)
                    if response.status_code == 200:
                        image_ext = mimetypes.guess_extension(response.headers["Content-Type"])
                        file_uuid = str(uuid.uuid4())
                        file_key = image_dir / f"{file_uuid}.{image_ext}"
                        get_storage().save(file_key, response.content)
                    else:
                        continue
                else:
                    image_ext = rel.target_ref.split(".")[-1]
                    # user uuid as file name
                    file_uuid = str(uuid.uuid4())
                    file_key = image_dir / f"{file_uuid}.{image_ext}"
                    get_storage().save(file_key, rel.target_part.blob)

                image_map[rel.target_part] = {
                    "text": f"![image_{image_count}](images/{file_uuid}.{image_ext})",
                    "image_name": f"image_{image_count}",
                    "image_data": encode_image_to_base64(file_key),
                }

        return image_map

    def _table_to_markdown(self, table, image_map):
        markdown = []
        rows = []
        # calculate the total number of columns
        total_cols = max(len(row.cells) for row in table.rows)

        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        rows.append(headers)
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

        for row in table.rows[1:]:
            row_cells = self._parse_row(row, image_map, total_cols)
            markdown.append("| " + " | ".join(row_cells) + " |")
            rows.append(row_cells)
        return "\n".join(markdown), rows

    def _parse_row(self, row, image_map, total_cols):
        # Initialize a row, all of which are empty by default
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            # make sure the col_index is not out of range
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            # if col_index is out of range the loop is jumped
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span if cell.grid_span else 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self, cell, image_map):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    image_part = paragraph.part.rels[image_id].target_part

                    if image_part in image_map:
                        image_link = image_map[image_part]["text"]
                        paragraph_content.append(image_link)
            else:
                paragraph_content.append(run.text)
        content = "".join(paragraph_content).strip()
        content = content.replace("\n", "<br>")
        return content

    def _parse_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id:
                        rel_target = run.part.rels[embed_id].target_ref
                        if rel_target in image_map:
                            paragraph_content.append(image_map[rel_target]["text"])
            if run.text.strip():
                paragraph_content.append(run.text.strip())
        return " ".join(paragraph_content) if paragraph_content else ""

    def parse_docx(self, docx_path):
        doc = DocxDocument(docx_path)

        content = []

        image_map = self._extract_images_from_docx(doc)

        hyperlinks_url = None
        url_pattern = re.compile(r"http://[^\s+]+//|https://[^\s+]+")
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and hyperlinks_url:
                    result = f"  [{run.text}]({hyperlinks_url})  "
                    run.text = result
                    hyperlinks_url = None
                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ET.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x_child is None:
                                continue
                            if x.tag.endswith("instrText"):
                                for i in url_pattern.findall(x.text):
                                    hyperlinks_url = str(i)
                    except Exception as e:
                        logger.error(e)

        def parse_paragraph(paragraph):
            paragraph_content = []
            chunks = []
            chunk_content = ""
            for run in paragraph.runs:
                if (
                    hasattr(run.element, "tag")
                    and isinstance(element.tag, str)
                    and run.element.tag.endswith("r")
                ):
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                image_part = doc.part.related_parts.get(embed_id)
                                if image_part in image_map:
                                    image_info = image_map[image_part]
                                    paragraph_content.append(image_info["text"])
                                    if chunk_content:
                                        chunks.append(Chunk(content=chunk_content))
                                        chunk_content = ""
                                    chunks.append(
                                        Chunk(
                                            content=image_info["text"],
                                            chunk_type="image",
                                            image_name=image_info["image_name"],
                                            image_content=image_info["image_data"],
                                        )
                                    )
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
                    chunk_content += run.text.strip()
            if chunk_content:
                chunks.append(Chunk(content=chunk_content))
            return "".join(paragraph_content) if paragraph_content else "", chunks

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        chunks = []
        for element in doc.element.body:
            if hasattr(element, "tag"):
                if isinstance(element.tag, str) and element.tag.endswith("p"):  # paragraph
                    para = paragraphs.pop(0)
                    parsed_paragraph, p_chunks = parse_paragraph(para)
                    if parsed_paragraph:
                        content.append(parsed_paragraph)
                        chunks.extend(p_chunks)
                elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # table
                    table = tables.pop(0)
                    markdown, rows = self._table_to_markdown(table, image_map)
                    content.append(markdown)
                    chunks.append(
                        Chunk(
                            content=markdown, chunk_type="table_csv", table_content=csv_dumps(rows)
                        )
                    )
        ret = Document(
            summary="\n".join(content),
            metadata={
                "num_paragraphs": len(content),
                "num_tables": len(tables),
                "num_images": len(image_map),
            },
        )
        ret.add_chunk(chunks)
        return ret
